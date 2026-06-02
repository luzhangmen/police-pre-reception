"""Generate the primary property-loss showcase and print before/after narrative."""

import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from docx import Document

from app.modules.document_generator import GENERATED_DIR, generate_pre_acceptance_document
from app.modules.narrative_polish import clean_colloquial_text, synthesize_incident_narrative
from scripts.showcase_states import SHOWCASE_CASES


def _extract_doc_text(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def main() -> None:
    state = SHOWCASE_CASES["property_loss"]
    out_dir = GENERATED_DIR / "showcase"
    out_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("【原始口语】\n", state.user_text)
    print("\n【去废话】\n", clean_colloquial_text(state.user_text))
    print("\n【事件经过】\n", synthesize_incident_narrative(state))
    print("=" * 60)

    path, name = generate_pre_acceptance_document(state, output_dir=out_dir)
    print(f"\nWord: {path}\n")
    print(_extract_doc_text(path))


if __name__ == "__main__":
    main()
