from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.core.state import CaseState
from app.main import app
from app.modules.document_content import build_document_bundle
from app.modules.document_generator import (
    build_document,
    build_placeholder_map,
    create_blank_template,
    generate_pre_acceptance_document,
)
from app.modules.document_styling import DOCUMENT_VERSION


def test_build_placeholder_map_uses_chinese_labels():
    state = CaseState(
        case_id="doc-test-001",
        user_text="测试描述",
        scenario="telecom_fraud",
        risk_level="high",
        completeness_score=0.5,
        slots={"reporter_name": "测试", "loss_amount": 100},
        missing_fields=["evidence"],
        knowledge_snippets=["先止付"],
        police_summary="摘要内容",
        suggested_next_steps=["保存证据"],
    )
    mapping = build_placeholder_map(state)
    assert mapping["{{case_id}}"] == "doc-test-001"
    assert mapping["{{risk_level}}"] == "高"
    assert "损失金额" in mapping["{{slots_block}}"]


def test_build_document_has_thirteen_sections(tmp_path: Path):
    state = CaseState(
        case_id="doc-test-002",
        user_text="图书馆丢了手机",
        scenario="property_loss",
        risk_level="high",
        slots={"reporter_name": "李同学", "lost_item": "手机", "suspected_theft": True},
        police_summary="遗失手机一份",
        key_facts=["地点：图书馆"],
    )
    doc = build_document(state)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "校园警务预受理信息单" in text
    assert DOCUMENT_VERSION in text
    assert "一、 报案原话" in text or "一、报案原话" in text.replace(" ", "")
    assert "十三、 民警备注区" in text or "十三、民警备注区" in text.replace(" ", "")
    assert "高风险预警" in text


def test_generate_pre_acceptance_document_writes_docx(tmp_path: Path):
    state = CaseState(
        case_id="doc-test-002",
        user_text="图书馆丢了手机",
        scenario="property_loss",
        slots={"reporter_name": "李同学", "lost_item": "手机"},
        police_summary="遗失手机一份",
    )
    output_path, filename = generate_pre_acceptance_document(state, output_dir=tmp_path)
    assert output_path.is_file()
    assert filename.endswith(".docx")

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "李同学" in full_text
    assert "{{" not in full_text


def test_create_blank_template(tmp_path: Path):
    target = tmp_path / "template.docx"
    path = create_blank_template(target)
    assert path.is_file()


def test_pre_acceptance_api_returns_metadata(tmp_path, monkeypatch):
    import app.main as main_app
    from app.modules import document_generator as dg

    monkeypatch.setattr(dg, "GENERATED_DIR", tmp_path)
    monkeypatch.setattr(main_app, "GENERATED_DIR", tmp_path)

    client = TestClient(app)
    payload = {
        "case_id": "api-doc-001",
        "user_text": "嗯那个我被骗了500元",
        "scenario": "telecom_fraud",
        "risk_level": "high",
        "slots": {"reporter_name": "测试用户", "still_transferring": True},
        "police_summary": "API 测试摘要",
    }
    create_resp = client.post("/api/v1/documents/pre-acceptance", json=payload)
    assert create_resp.status_code == 200
    body = create_resp.json()
    assert body["metadata"]["document_version"] == "2.0"
    assert "narrative_preview" in body["metadata"]

    download_resp = client.get(body["download_url"])
    assert download_resp.status_code == 200


def test_document_bundle_high_risk_flags():
    state = CaseState(
        case_id="x",
        user_text="test",
        scenario="telecom_fraud",
        slots={"still_transferring": True, "still_contacting": True},
    )
    bundle = build_document_bundle(state)
    assert len(bundle["high_risk_flags"]) >= 1
